import os
import re
import markdown
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Absolute paths
BASE_DIR = r"c:\Users\Aashutosh\Desktop\SHL-Task\talentroute-ai"
MD_PATH = os.path.join(BASE_DIR, "Approach_Document.md")
PDF_PATH = os.path.join(BASE_DIR, "Approach_Document.pdf")
DOCX_PATH = os.path.join(BASE_DIR, "Approach_Document.docx")
SUBMISSION_DOCX_PATH = os.path.join(BASE_DIR, "IIT-KGP-Aashutosh-Joshi-SHL-Submission.docx")

def convert_to_pdf():
    print("Starting PDF Conversion...")
    # Read Markdown content
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_content = f.read()

    base_dir_unix = BASE_DIR.replace('\\', '/')
    absolute_img_prefix = base_dir_unix + '/'
    md_content = md_content.replace('src="image/', f'src="{absolute_img_prefix}image/')

    # Convert MD to HTML with extensions
    html_body = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'toc'])

    # Premium, academic CSS style sheet for PDF
    css = """
    @page {
        size: letter;
        margin: 0.8in;
    }
    body {
        font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
        color: #333333;
        line-height: 1.5;
        font-size: 10.5pt;
    }
    h1 {
        font-size: 22pt;
        color: #1a365d;
        border-bottom: 2px solid #2b5c8f;
        padding-bottom: 8px;
        margin-top: 30px;
        margin-bottom: 15px;
    }
    h2 {
        font-size: 15pt;
        color: #2b5c8f;
        margin-top: 25px;
        margin-bottom: 12px;
        border-bottom: 1px solid #dddddd;
        padding-bottom: 4px;
    }
    h3 {
        font-size: 12pt;
        color: #319795;
        margin-top: 20px;
        margin-bottom: 10px;
    }
    h4 {
        font-size: 11pt;
        color: #2d3748;
        font-weight: bold;
    }
    p {
        margin-bottom: 12px;
        text-align: justify;
    }
    ul, ol {
        margin-bottom: 12px;
        padding-left: 20px;
    }
    li {
        margin-bottom: 4px;
    }
    table {
        width: 100%;
        border-collapse: collapse;
        margin-top: 15px;
        margin-bottom: 20px;
        font-size: 9.5pt;
    }
    th {
        background-color: #f7fafc;
        color: #2d3748;
        font-weight: bold;
        border-bottom: 2px solid #cbd5e0;
        padding: 8px 10px;
        text-align: left;
    }
    td {
        border-bottom: 1px solid #e2e8f0;
        padding: 8px 10px;
        text-align: left;
    }
    tr:nth-child(even) td {
        background-color: #fcfcfc;
    }
    code {
        font-family: 'Courier New', Courier, monospace;
        background-color: #edf2f7;
        padding: 2px 4px;
        font-size: 9pt;
        border-radius: 3px;
        color: #2d3748;
    }
    pre {
        background-color: #edf2f7;
        padding: 12px;
        border-radius: 5px;
        margin-bottom: 15px;
        font-family: 'Courier New', Courier, monospace;
        font-size: 8.5pt;
        line-height: 1.3;
    }
    pre code {
        background-color: transparent;
        padding: 0;
        color: inherit;
    }
    img {
        display: block;
        margin-left: auto;
        margin-right: auto;
        max-width: 90%;
        margin-top: 15px;
        margin-bottom: 10px;
    }
    p.center-em {
        text-align: center;
        font-style: italic;
        color: #555555;
        font-size: 9pt;
        margin-top: 5px;
        margin-bottom: 20px;
    }
    hr {
        border: 0;
        border-top: 1px solid #e2e8f0;
        margin: 25px 0;
    }
    """

    # Wrap the body in standard HTML boilerplate
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            {css}
        </style>
    </head>
    <body>
        {html_body}
    </body>
    </html>
    """

    # Replace figure wrapper <p align="center"> to look beautiful in PDF
    full_html = re.sub(
        r'<p align="center">\s*<img([^>]+)>\s*<br>\s*<em>(.*?)</em>\s*</p>',
        r'<div style="text-align: center; margin: 20px 0;"><img\1><p class="center-em">\2</p></div>',
        full_html,
        flags=re.DOTALL
    )

    # Write output PDF
    with open(PDF_PATH, "wb") as f_pdf:
        pisa_status = pisa.CreatePDF(full_html, dest=f_pdf)
    
    if pisa_status.err:
        print("Error during PDF generation.")
    else:
        print(f"Successfully generated: {PDF_PATH}")


def set_cell_background(cell, color_hex):
    """Set shading color for table cells."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def convert_to_docx():
    print("Starting Word DOCX Conversion...")
    doc = Document()

    # Premium style definitions
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base text styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Helper function to parse markdown tables
    def parse_md_table(table_lines):
        data = []
        for line in table_lines:
            # Strip outer pipes
            cleaned = line.strip().strip('|')
            cells = [c.strip() for c in cleaned.split('|')]
            data.append(cells)
        return data

    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_table = False
    table_lines = []
    in_code_block = False

    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        # Handle Code Blocks
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            idx += 1
            continue

        if in_code_block:
            idx += 1
            continue

        # Handle Tables
        if stripped.startswith('|'):
            in_table = True
            table_lines.append(line)
            idx += 1
            continue
        elif in_table:
            # End of table
            in_table = False
            table_data = parse_md_table(table_lines)
            table_lines = []
            
            # Construct table in Word doc (skipping separation lines e.g. :---)
            cleaned_data = []
            for row in table_data:
                if len(row) > 0 and all(cell.startswith(':') or cell.startswith('-') or cell == '' for cell in row):
                    continue # Skip spacer row
                cleaned_data.append(row)

            if len(cleaned_data) > 0:
                cols_count = len(cleaned_data[0])
                docx_table = doc.add_table(rows=len(cleaned_data), cols=cols_count)
                docx_table.autofit = True
                
                # Populate cells
                for r_idx, row in enumerate(cleaned_data):
                    for c_idx, val in enumerate(row):
                        if c_idx < len(docx_table.rows[r_idx].cells):
                            # Clean cell content (strip strong marks e.g. **)
                            cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', val).strip()
                            cell_text = re.sub(r'\`(.*?)\`', r'\1', cell_text).strip()
                            cell = docx_table.rows[r_idx].cells[c_idx]
                            cell.text = cell_text
                            
                            # Styling Header row
                            if r_idx == 0:
                                set_cell_background(cell, "F2F2F2")
                                cell.paragraphs[0].runs[0].font.bold = True
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x11, 0x11, 0x11)
                
                doc.add_paragraph() # Add margin space after table
            idx += 1
            continue

        # Handle Images
        if "<img" in stripped or "![" in stripped:
            img_name = None
            if "embedding_comparison" in stripped:
                img_name = "embedding_comparison.png"
            elif "threshold_sweep" in stripped:
                img_name = "threshold_sweep.png"
            elif "architecture_flow" in stripped:
                img_name = "architecture_flow.png"
            
            if img_name:
                img_file_path = os.path.join(BASE_DIR, "image", img_name)
                if os.path.exists(img_file_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run()
                    r.add_picture(img_file_path, width=Inches(5.0))
                    
                    # Extract figure label if it is next line
                    if idx + 1 < len(lines) and "Figure" in lines[idx+1] or "<em>" in lines[idx+1]:
                        fig_line = lines[idx+1].strip()
                        fig_text = re.sub(r'<.*?>', '', fig_line).strip() # Remove tags
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r_cap = p_cap.add_run(fig_text)
                        r_cap.font.italic = True
                        r_cap.font.size = Pt(9.5)
                        r_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                        idx += 2
                        continue
            idx += 1
            continue

        # Handle Headings
        if stripped.startswith('#'):
            h_match = re.match(r'^(#{1,4})\s+(.*)$', stripped)
            if h_match:
                level = len(h_match.group(1))
                heading_text = h_match.group(2).strip()
                # Clean strong/code formatting inside headers
                heading_text = re.sub(r'\*\*(.*?)\*\*', r'\1', heading_text)
                heading_text = re.sub(r'\`(.*?)\`', r'\1', heading_text)
                
                h = doc.add_heading(heading_text, level=level)
                
                # Custom heading coloring
                run = h.runs[0] if len(h.runs) > 0 else h.add_run(heading_text)
                if level == 1:
                    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
                    run.font.size = Pt(18)
                elif level == 2:
                    run.font.color.rgb = RGBColor(0x2B, 0x5C, 0x8F)
                    run.font.size = Pt(14)
                else:
                    run.font.color.rgb = RGBColor(0x31, 0x97, 0x95)
                    run.font.size = Pt(12)
            idx += 1
            continue

        # Handle List Items
        if stripped.startswith('*') or stripped.startswith('-'):
            list_text = re.sub(r'^[\*\-\+]\s+', '', stripped)
            list_text = re.sub(r'\*\*(.*?)\*\*', r'\1', list_text)
            list_text = re.sub(r'\`(.*?)\`', r'\1', list_text)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run(list_text)
            idx += 1
            continue

        # Handle Empty Lines
        if stripped == '':
            idx += 1
            continue

        # Regular Paragraph
        cleaned_text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
        cleaned_text = re.sub(r'\`(.*?)\`', r'\1', cleaned_text)
        
        # Skip isolated figure lines/center wrapper text if already processed
        if cleaned_text.startswith('<p') or cleaned_text.startswith('</p') or cleaned_text.startswith('<br'):
            idx += 1
            continue
            
        doc.add_paragraph(cleaned_text)
        idx += 1

    # Save DOCX
    try:
        doc.save(DOCX_PATH)
        print(f"Successfully generated: {DOCX_PATH}")
    except PermissionError:
        fallback_path = DOCX_PATH.replace(".docx", "_temp.docx")
        doc.save(fallback_path)
        print(f"Warning: {DOCX_PATH} was locked (likely open in Word). Saved to fallback: {fallback_path}")

    try:
        doc.save(SUBMISSION_DOCX_PATH)
        print(f"Successfully generated: {SUBMISSION_DOCX_PATH}")
    except PermissionError:
        fallback_path = SUBMISSION_DOCX_PATH.replace(".docx", "_temp.docx")
        doc.save(fallback_path)
        print(f"Warning: {SUBMISSION_DOCX_PATH} was locked (likely open in Word). Saved to fallback: {fallback_path}")

if __name__ == "__main__":
    convert_to_pdf()
    convert_to_docx()
